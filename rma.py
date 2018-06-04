# -*- coding: utf-8 -*-
"""
Created on Fri Feb  2 14:49:37 2018

@author: hmamin
"""

import sys
import os
import time
import glob
import re
import pyperclip
from pprint import pprint
import docx
from docx.shared import Pt
import comtypes.client as com
from oauth2client.service_account import ServiceAccountCredentials
import gspread


class Swap():
    '''One swap row from RMA tracker.'''

    def __init__(self):
        '''Initialize Swap() instance.'''
        self.sn = None
        self.row = None
        self.date = None
        self.site = None
        self.tower = None
        self.issue = None
        self.is_install = None

    def rma_pop(self, sheet, c):
        '''After finding a swap row in RMA, fill in other field details.'''
        self.sn = c.value
        self.row = c.row
        self.date = sheet.cell(c.row, 2).value
        self.issue = sheet.cell(c.row, 10).value
        self.tower = 'Tower '
        self.is_install = False
        
        # Parse RMA sheet "Comments" cell.                       
        comments = sheet.cell(c.row, 13).value
        com_split = comments.lower().split('\n')                    
        if len(com_split) == 2:
             self.site, self.tower = com_split                    
        else:
            self.site = re.sub('\s*((t)|(tower))\s*\d+', '', comments.lower())
            matches = re.search('(?<=tower)\s?\d', comments)   
            if matches is not None:
                self.tower += matches.group().strip()
                
    def mdt_pop(self, sheet, c):
        '''Pass in worksheet "sheet" and cell "c" to populate first swap 
        instance (the installation).'''
        self.sn = serial
        self.row = c.row
        self.date = sheet.cell(c.row, 9).value
        self.site = sheet.cell(c.row, 1).value
        self.is_install = True
        
        tower_match = re.search('\d(?=:)', c.value)  
        if tower_match is None:                    
            self.tower = 'Tower 1'
        else:
            self.tower = 'Tower ' + tower_match.group()
  
                    
def pause(action='continue'):
    '''Pause program until user hits ENTER key.'''
    input('\nPress ENTER to ' + action + ': ')


def authorize_drive(file_name):    
    '''Use creds to create a client to interact with the Google Drive API.'''
    scope = ['https://spreadsheets.google.com/feeds']
    creds = ServiceAccountCredentials.from_json_keyfile_name(file_name, scope)
    client = gspread.authorize(creds)
    return client


# Authorize access to Google Drive
months = {'January', 'February', 'March', 'April', 'May', 'June', 'July', \
'August', 'September', 'October', 'November', 'December'}
client = authorize_drive('client_secret.json')
serial = pyperclip.paste().strip()
print('Serial #:', serial)
swapDate = ''
installDate = ''

# Find a workbook by name and open the first sheet
MDT = client.open('40-000055 Master Deployment Tracking').sheet1
RMA = client.open('40-000143 RMA & Servicing Tracker').sheet1
print('Searching for ' + str(serial) + '...') 

# create empty list to fill with sites, then find latest        
swaps = []
                                      
# Iterate through sn matches, turning each row into a Swap instance in list of swaps. 
print('\n' + '-' * 20 + '\nRMA Search\n' + '-' * 20)
RMAcells = RMA.findall(serial)     
for c in RMAcells:
    if c.col == 6:
        current_swap = Swap()
        current_swap.rma_pop(RMA, c)
        print('\nRow ' + str(current_swap.row) + ': ' + current_swap.site)
        print(current_swap.date + ' [' + current_swap.sn + ']')
        swaps.append(current_swap)            

# Iterate through MDT matches (hopefully only 1).        
print('\n' + '-' * 20 + '\nMDT Search\n' + '-' * 20)
MDTcells = MDT.range('T17:T{}'.format(MDT.row_count))
for c in MDTcells:
    if serial in c.value:
        install_swap = Swap()
        install_swap.mdt_pop(MDT, c)
        swaps.insert(0, install_swap)
        print('\nRow ' + str(install_swap.row) + ': ' + install_swap.site)
        print(install_swap.date + ' [' + serial + ']')
print('\n' + '-' * 25 + '\nSearch complete\n' + '-'*25)  
if len(swaps) == 0:
    sys.exit('No record of this serial #.')
pprint(swaps[-1].__dict__)

# Start populating info for new row in RMA tracker
new_row = {
'hidden': '',
'date': time.strftime('%m/%d/%Y'),
'agiloft': '',
'part': 'Inverter',
'old_serial': serial,
'new_serial': '',
'vendor': 'Ideal Power Converters',
'rmaNum': '',
'fedex': '',
'problem': '',
'status': '',
'next': '',
'comments': ''
}
new_row['problem'] = input('Paste summary of fault from Agiloft:\n')

# Set backup value for current site in case not listed on MDT    
try:
    current_site = swaps[-1].site.title()
except Exception:
    current_site = 'SITE_NAME_PLACEHOLDER'
    
# Create format for new word file
today = time.strftime('%m.%d.%Y')
name_format = r'C:\Users\hmamin\Desktop\RMA\RMARequest_Rev E_FOR-' +\
        serial[-4:] + '_' + current_site + '_' + today + '.docx'
       
# Copy most recently edited word doc in RMA folder; backup option in case of corruption
oldCopy = max(glob.iglob(r'C:\Users\hmamin\Desktop\RMA\*.docx'), 
              key=os.path.getmtime)
backup_old_copy = r'C:\Users\hmamin\Desktop\RMA\RMARequest_Rev E_FOR-0021_Tulare USD-Central Kitchen_09.26.2017.docx'

# Create copy of file and populate table.
try:
    doc = docx.Document(oldCopy)
    print('\nOld file copied from: ', oldCopy)
except Exception as e:
    doc = docx.Document(backup_old_copy)
    print('\nOld file copied from: ', backup_old_copy)

# Loop through table rows. Put sn in row 4, date in row 6, site/tower in row 8.
table = doc.tables[0]
for row in range(4, 10, 2):
    c_cell = table.cell(0, row)    
    if row != 8:
        c_cell.paragraphs[0]._p.clear()
        if row == 4:
            c_cell.paragraphs[0].add_run('Serial Number: ').bold = True
            c_cell.paragraphs[0].add_run(serial) 
        else:
            c_cell.paragraphs[0].add_run('Installation Date: ').bold = True
            c_cell.paragraphs[0].add_run(swaps[-1].date)
        c_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    else:
        c_cell.paragraphs[1].text = current_site
        c_cell.paragraphs[2].text = swaps[-1].tower
doc.paragraphs[5].text = new_row['problem']
doc.save(name_format)
print('\nNew file generated: ', name_format)

# Double check word doc, then hit ENTER to resume program
pause('generate PDF and update RMA tracker')  
word_copy = max(glob.iglob(r'C:\Users\hmamin\Desktop\RMA\*.docx'), 
                key=os.path.getmtime)
print('\nOld file copied from: ', word_copy)

# Copy word file to pdf.
word = com.CreateObject('Word.Application')
comDoc = word.Documents.Open(word_copy)
comDoc.SaveAs(word_copy[:-5] + '.pdf', FileFormat=17)
comDoc.Close()
word.Quit()

# Get text from word doc table.
table_con = []
tables = doc.tables
for row in tables[0].rows:
    table_con.append(row.cells[0].text)

# User adds in ticket # and new serial # if known.
new_row['agiloft'] = input('\nTicket # (Hit ENTER if n/a): ')
new_row['new_serial'] = input('\nNew serial # (Hit ENTER if unknown): ')   
new_row['comments'] = table_con[8].replace('Installation Location (Name of\
 Location, Address/City/State, Tower Site or Marker):\n', '')
new_row['problem'] = doc.paragraphs[5].text     #make sure no extra blank line? 
pprint(new_row)

# Copy message to email to IPWR
email = f'Lee, \n\nThe RMA form for {current_site} is attached below. \
When you get a chance, please send a shipping label for the old unit.\
\n\nBest,\n\nHarrison Mamin'
pyperclip.copy(email)

# Add new row to RMA
RMA.append_row(new_row.values())